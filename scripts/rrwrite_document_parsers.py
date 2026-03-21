#!/usr/bin/env python3
"""
Document Parsers - Parse journal requirement documents from various formats.

This module provides parsers for extracting journal submission requirements
from PDF, DOCX, HTML, and YAML sources.
"""

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Any
import yaml

# Optional imports (will gracefully degrade if not available)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import requests
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


class ParsingError(Exception):
    """Exception raised when document parsing fails."""
    pass


class BaseDocumentParser(ABC):
    """Abstract base class for document parsers."""

    def __init__(self, source: str):
        """
        Initialize parser.

        Args:
            source: Path to file or URL
        """
        self.source = source
        self.raw_text = ""
        self.structured_data = {}

    @abstractmethod
    def parse(self) -> Dict[str, Any]:
        """
        Parse the document and extract structured data.

        Returns:
            Dictionary containing parsed data

        Raises:
            ParsingError: If parsing fails
        """
        pass

    def get_text(self) -> str:
        """Return raw extracted text."""
        return self.raw_text


class PDFParser(BaseDocumentParser):
    """Parser for PDF documents."""

    def __init__(self, pdf_path: str):
        """
        Initialize PDF parser.

        Args:
            pdf_path: Path to PDF file
        """
        if not PDF_AVAILABLE:
            raise ParsingError("PyPDF2 not installed. Install with: pip install PyPDF2")

        super().__init__(pdf_path)
        self.pdf_path = Path(pdf_path)

        if not self.pdf_path.exists():
            raise ParsingError(f"PDF file not found: {pdf_path}")

    def parse(self) -> Dict[str, Any]:
        """
        Parse PDF and extract text.

        Returns:
            Dictionary with 'text' and 'pages' keys
        """
        try:
            pages_text = []
            with open(self.pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    pages_text.append(text)

            self.raw_text = "\n".join(pages_text)
            self.structured_data = {
                'text': self.raw_text,
                'pages': pages_text,
                'num_pages': len(pages_text)
            }

            return self.structured_data

        except Exception as e:
            raise ParsingError(f"Failed to parse PDF: {e}")


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX documents."""

    def __init__(self, docx_path: str):
        """
        Initialize DOCX parser.

        Args:
            docx_path: Path to DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ParsingError("python-docx not installed. Install with: pip install python-docx")

        super().__init__(docx_path)
        self.docx_path = Path(docx_path)

        if not self.docx_path.exists():
            raise ParsingError(f"DOCX file not found: {docx_path}")

    def parse(self) -> Dict[str, Any]:
        """
        Parse DOCX and extract text with structure.

        Returns:
            Dictionary with 'text', 'paragraphs', and 'tables' keys
        """
        try:
            doc = Document(self.docx_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else None
                    })

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            self.raw_text = "\n".join([p['text'] for p in paragraphs])
            self.structured_data = {
                'text': self.raw_text,
                'paragraphs': paragraphs,
                'tables': tables
            }

            return self.structured_data

        except Exception as e:
            raise ParsingError(f"Failed to parse DOCX: {e}")


class HTMLParser(BaseDocumentParser):
    """Parser for HTML documents (web pages)."""

    def __init__(self, url: str, timeout: int = 30):
        """
        Initialize HTML parser.

        Args:
            url: URL to fetch
            timeout: Request timeout in seconds
        """
        if not HTML_AVAILABLE:
            raise ParsingError(
                "BeautifulSoup4 and requests not installed. "
                "Install with: pip install beautifulsoup4 requests lxml"
            )

        super().__init__(url)
        self.url = url
        self.timeout = timeout

    def parse(self) -> Dict[str, Any]:
        """
        Fetch and parse HTML content.

        Returns:
            Dictionary with 'text', 'soup', and 'sections' keys
        """
        try:
            # Fetch HTML
            response = requests.get(self.url, timeout=self.timeout)
            response.raise_for_status()

            # Parse with BeautifulSoup
            soup = BeautifulSoup(response.content, 'lxml')

            # Remove script and style elements
            for element in soup(['script', 'style', 'nav', 'footer']):
                element.decompose()

            # Extract text
            self.raw_text = soup.get_text(separator='\n', strip=True)

            # Extract structured sections
            sections = self._extract_sections(soup)

            # Extract tables
            tables = self._extract_tables(soup)

            self.structured_data = {
                'text': self.raw_text,
                'soup': soup,
                'sections': sections,
                'tables': tables,
                'url': self.url
            }

            return self.structured_data

        except requests.RequestException as e:
            raise ParsingError(f"Failed to fetch URL: {e}")
        except Exception as e:
            raise ParsingError(f"Failed to parse HTML: {e}")

    def _extract_sections(self, soup: Any) -> List[Dict[str, str]]:
        """Extract sections based on heading tags."""
        sections = []

        # Find all headings
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4']):
            section = {
                'level': int(heading.name[1]),
                'title': heading.get_text(strip=True),
                'content': []
            }

            # Get content until next heading
            for sibling in heading.find_next_siblings():
                if sibling.name in ['h1', 'h2', 'h3', 'h4']:
                    break
                if sibling.get_text(strip=True):
                    section['content'].append(sibling.get_text(strip=True))

            sections.append(section)

        return sections

    def _extract_tables(self, soup: Any) -> List[List[List[str]]]:
        """Extract tables from HTML."""
        tables = []

        for table in soup.find_all('table'):
            table_data = []
            for row in table.find_all('tr'):
                row_data = []
                for cell in row.find_all(['td', 'th']):
                    row_data.append(cell.get_text(strip=True))
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return tables


class YAMLConverter(BaseDocumentParser):
    """Converter for existing YAML journal entries."""

    def __init__(self, yaml_path: str, journal_name: str):
        """
        Initialize YAML converter.

        Args:
            yaml_path: Path to journal_guidelines.yaml
            journal_name: Name of journal to extract
        """
        super().__init__(yaml_path)
        self.yaml_path = Path(yaml_path)
        self.journal_name = journal_name

        if not self.yaml_path.exists():
            raise ParsingError(f"YAML file not found: {yaml_path}")

    def parse(self) -> Dict[str, Any]:
        """
        Load and extract journal data from YAML.

        Returns:
            Dictionary containing journal data
        """
        try:
            with open(self.yaml_path, 'r') as f:
                data = yaml.safe_load(f)

            # Find journal entry
            # Journals are stored as a dict with keys like 'bioinformatics', 'nature_methods', etc.
            journal_data = None
            if isinstance(data, dict) and 'journals' in data:
                journals = data['journals']

                # Try exact key match first (normalize journal name to key format)
                journal_key = self.journal_name.lower().replace(' ', '_')
                if journal_key in journals:
                    journal_data = journals[journal_key]
                else:
                    # Try matching full_name field
                    for key, journal_info in journals.items():
                        if isinstance(journal_info, dict):
                            full_name = journal_info.get('full_name', '')
                            if self.journal_name.lower() in full_name.lower():
                                journal_data = journal_info
                                break

            if not journal_data:
                raise ParsingError(
                    f"Journal '{self.journal_name}' not found in YAML file"
                )

            self.structured_data = journal_data
            self.raw_text = str(journal_data)

            return self.structured_data

        except yaml.YAMLError as e:
            raise ParsingError(f"Failed to parse YAML: {e}")
        except Exception as e:
            raise ParsingError(f"Failed to load YAML: {e}")


def create_parser(source: str, source_type: str, **kwargs) -> BaseDocumentParser:
    """
    Factory function to create appropriate parser.

    Args:
        source: Path to file or URL
        source_type: Type of source ('pdf', 'docx', 'html', 'yaml')
        **kwargs: Additional arguments for specific parsers

    Returns:
        Appropriate parser instance

    Raises:
        ParsingError: If source_type is unknown
    """
    source_type = source_type.lower()

    if source_type == 'pdf':
        return PDFParser(source)
    elif source_type == 'docx':
        return DOCXParser(source)
    elif source_type in ['html', 'url']:
        return HTMLParser(source, **kwargs)
    elif source_type == 'yaml':
        journal_name = kwargs.get('journal_name')
        if not journal_name:
            raise ParsingError("journal_name required for YAML parser")
        return YAMLConverter(source, journal_name)
    else:
        raise ParsingError(f"Unknown source type: {source_type}")


if __name__ == '__main__':
    # Test parsers
    import sys

    if len(sys.argv) < 3:
        print("Usage: python rrwrite_document_parsers.py <source_type> <source>")
        print("  source_type: pdf, docx, html, yaml")
        print("  source: path to file or URL")
        sys.exit(1)

    source_type = sys.argv[1]
    source = sys.argv[2]

    try:
        parser = create_parser(source, source_type)
        data = parser.parse()
        print(f"Successfully parsed {source_type.upper()}")
        print(f"Extracted {len(parser.get_text())} characters of text")
    except ParsingError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
