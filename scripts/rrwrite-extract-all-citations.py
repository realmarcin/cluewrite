#!/usr/bin/env python3
"""
Enhanced citation extractor that captures ALL citation formats:
- Paperpile hyperlinks
- Bracketed citations
- Plain text citations (Author et al. Year)
- Numbered citations
- References section (if present)

Usage:
    python scripts/rrwrite-extract-all-citations.py \
        --docx manuscript.docx \
        --output all_citations.json \
        --stats
"""

import argparse
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Set, Tuple
from collections import defaultdict

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Error: python-docx not installed")
    exit(1)


def extract_paperpile_citations(text: str) -> List[Dict]:
    """Extract Paperpile hyperlink citations"""
    pattern = r'\[(.*?)\]\(https://paperpile\.com/c/([^/]+)/([^\)]+)\)'
    citations = []
    seen = set()

    for match in re.finditer(pattern, text):
        display_text = match.group(1)
        project_code = match.group(2)
        paperpile_code = match.group(3)

        author, year = parse_author_year(display_text)
        citation_key = f"{project_code}/{paperpile_code}"

        if citation_key not in seen:
            citations.append({
                'type': 'paperpile',
                'display_text': display_text,
                'author': author,
                'year': year,
                'source': 'hyperlink'
            })
            seen.add(citation_key)

    return citations


def extract_plain_text_citations(text: str) -> List[Dict]:
    """
    Extract plain text citations like:
    - Author et al. (2024)
    - Author and Author (2024)
    - Author et al. 2024
    - (Author et al., 2024)
    """
    citations = []
    seen = set()

    # Pattern 1: Author et al. (YEAR) or Author et al. YEAR
    pattern1 = r'\b([A-Z][a-z]+(?:\s+(?:and|&)\s+[A-Z][a-z]+)?)\s+et\s+al\.?\s*[\(,]?\s*(\d{4})\)?'

    # Pattern 2: Author and Author (YEAR)
    pattern2 = r'\b([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)\s*[\(,]?\s*(\d{4})\)?'

    # Pattern 3: Single author (YEAR)
    pattern3 = r'\b([A-Z][a-z]+)\s*[\(,]\s*(\d{4})\)'

    # Pattern 4: (Author et al., YEAR)
    pattern4 = r'\(([A-Z][a-z]+(?:\s+(?:and|&)\s+[A-Z][a-z]+)?)\s+et\s+al\.?,?\s*(\d{4})\)'

    patterns = [
        (pattern1, lambda m: (m.group(1), m.group(2))),
        (pattern2, lambda m: (f"{m.group(1)} and {m.group(2)}", m.group(3))),
        (pattern3, lambda m: (m.group(1), m.group(2))),
        (pattern4, lambda m: (m.group(1), m.group(2)))
    ]

    for pattern, extract_fn in patterns:
        for match in re.finditer(pattern, text):
            try:
                author, year = extract_fn(match)

                # Clean up author name
                author = author.strip()

                # Create citation key
                citation_key = f"{author.replace(' ', '')}_{year}"

                if citation_key not in seen:
                    citations.append({
                        'type': 'plain_text',
                        'author': author,
                        'year': year,
                        'display_text': match.group(0),
                        'source': 'inline'
                    })
                    seen.add(citation_key)
            except Exception:
                continue

    return citations


def extract_bracketed_citations(text: str) -> List[Dict]:
    """Extract bracketed citations like [Author2024]"""
    pattern = r'\[([A-Z][a-zA-Z]+\d{4}[a-z]?)\]'
    citations = []
    seen = set()

    for match in re.finditer(pattern, text):
        citation_key = match.group(1)

        if citation_key not in seen:
            year_match = re.search(r'(\d{4})', citation_key)
            author = citation_key[:year_match.start()] if year_match else citation_key
            year = year_match.group(1) if year_match else None

            citations.append({
                'type': 'bracketed',
                'citation_key': citation_key,
                'author': author,
                'year': year,
                'source': 'inline'
            })
            seen.add(citation_key)

    return citations


def extract_references_section(doc: Document) -> Tuple[List[Dict], int, int]:
    """
    Extract citations from References/Bibliography section if present.

    Returns:
        Tuple of (citations, start_index, end_index)
    """
    citations = []

    # Find references section
    ref_start = -1
    ref_end = -1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()

        # Check for references header
        if text in ['references', 'bibliography', 'citations', 'works cited', 'literature cited']:
            ref_start = i
            continue

        # Check for end of references
        if ref_start >= 0 and text:
            if any(keyword in text for keyword in ['appendix', 'supplementary', 'acknowledgment', 'supporting information']):
                ref_end = i
                break

    if ref_start < 0:
        return citations, -1, -1

    if ref_end < 0:
        ref_end = len(doc.paragraphs)

    # Parse reference entries
    for i in range(ref_start + 1, ref_end):
        para_text = doc.paragraphs[i].text.strip()

        if not para_text:
            continue

        # Try to extract author and year from reference
        # Common formats:
        # Author, A., et al. (2024). Title. Journal.
        # Author A, Author B (2024) Title. Journal.

        author_match = re.search(r'^([A-Z][a-z]+(?:,\s*[A-Z]\.?)?)', para_text)
        year_match = re.search(r'\((\d{4})\)', para_text)

        if author_match and year_match:
            author = author_match.group(1).replace(',', '').strip()
            year = year_match.group(1)

            citations.append({
                'type': 'reference_entry',
                'author': author,
                'year': year,
                'full_text': para_text,
                'source': 'references_section'
            })

    return citations, ref_start, ref_end


def parse_author_year(text: str) -> Tuple[str, str]:
    """Parse author and year from display text"""
    text = text.strip().strip('()')

    year_match = re.search(r'\b(19\d{2}|20\d{2})\b', text)
    year = year_match.group(1) if year_match else None

    author_match = re.search(r'^([A-Z][a-z]+)', text)
    author = author_match.group(1) if author_match else None

    return author, year


def deduplicate_citations(citations: List[Dict]) -> List[Dict]:
    """Remove duplicates, prioritizing Paperpile > References > Plain text"""
    seen_keys = {}
    priority = {'paperpile': 1, 'reference_entry': 2, 'bracketed': 3, 'plain_text': 4}

    for cit in citations:
        # Generate key (handle None author/year)
        author = (cit.get('author') or '').replace(' ', '').replace('and', '').lower()
        year = cit.get('year') or ''
        key = f"{author}_{year}"

        # Check if we've seen this citation
        if key in seen_keys:
            # Keep higher priority version
            existing = seen_keys[key]
            if priority.get(cit['type'], 99) < priority.get(existing['type'], 99):
                seen_keys[key] = cit
        else:
            seen_keys[key] = cit

    return list(seen_keys.values())


def generate_statistics(citations: List[Dict]) -> Dict:
    """Generate statistics about citations"""
    by_type = defaultdict(int)
    by_source = defaultdict(int)
    by_year = defaultdict(int)
    by_author = defaultdict(int)

    for cit in citations:
        by_type[cit['type']] += 1
        by_source[cit.get('source', 'unknown')] += 1

        year = cit.get('year')
        if year:
            by_year[year] += 1

        author = cit.get('author')
        if author:
            # Extract first author
            first_author = author.split()[0].replace(',', '')
            by_author[first_author] += 1

    return {
        'total_citations': len(citations),
        'by_type': dict(by_type),
        'by_source': dict(by_source),
        'by_year': dict(sorted(by_year.items(), reverse=True)),
        'by_author': dict(sorted(by_author.items(), key=lambda x: x[1], reverse=True)[:10]),
        'year_range': {
            'earliest': min(by_year.keys()) if by_year else None,
            'latest': max(by_year.keys()) if by_year else None
        }
    }


def main():
    parser = argparse.ArgumentParser(
        description='Extract ALL citations from DOCX (Paperpile, plain text, references)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--docx', type=Path, required=True, help='Input DOCX file')
    parser.add_argument('--output', type=Path, required=True, help='Output JSON file')
    parser.add_argument('--stats', action='store_true', help='Print statistics')
    parser.add_argument('--verbose', action='store_true', help='Verbose output')

    args = parser.parse_args()

    if not args.docx.exists():
        print(f"Error: DOCX not found: {args.docx}")
        return 1

    # Read DOCX
    if args.verbose:
        print(f"Reading DOCX: {args.docx}")

    doc = Document(args.docx)
    full_text = '\n'.join([para.text for para in doc.paragraphs])

    # Extract all citation types
    all_citations = []

    # 1. Paperpile hyperlinks
    paperpile_cites = extract_paperpile_citations(full_text)
    all_citations.extend(paperpile_cites)
    if args.verbose:
        print(f"  Paperpile citations: {len(paperpile_cites)}")

    # 2. Plain text citations
    plain_text_cites = extract_plain_text_citations(full_text)
    all_citations.extend(plain_text_cites)
    if args.verbose:
        print(f"  Plain text citations: {len(plain_text_cites)}")

    # 3. Bracketed citations
    bracketed_cites = extract_bracketed_citations(full_text)
    all_citations.extend(bracketed_cites)
    if args.verbose:
        print(f"  Bracketed citations: {len(bracketed_cites)}")

    # 4. References section
    ref_cites, ref_start, ref_end = extract_references_section(doc)
    all_citations.extend(ref_cites)
    if args.verbose:
        if ref_start >= 0:
            print(f"  References section: paragraphs {ref_start}-{ref_end} ({len(ref_cites)} entries)")
        else:
            print(f"  References section: not found")

    # Deduplicate
    citations = deduplicate_citations(all_citations)

    if args.verbose:
        print(f"  Total unique citations: {len(citations)}")

    # Generate statistics
    stats = generate_statistics(citations)

    # Build output
    output_data = {
        'docx_file': str(args.docx),
        'extracted_at': datetime.now().isoformat(),
        'statistics': stats,
        'citations': citations
    }

    # Write output
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2)

    print(f"✓ Extracted {len(citations)} unique citations")
    print(f"✓ Output: {args.output}")

    # Print statistics
    if args.stats:
        print("\n=== Citation Statistics ===")
        print(f"Total citations: {stats['total_citations']}")

        print(f"\nBy type:")
        for cit_type, count in stats['by_type'].items():
            print(f"  {cit_type}: {count}")

        print(f"\nBy source:")
        for source, count in stats['by_source'].items():
            print(f"  {source}: {count}")

        if stats['year_range']['earliest']:
            print(f"\nYear range: {stats['year_range']['earliest']} - {stats['year_range']['latest']}")

        print(f"\nTop years:")
        for year, count in list(stats['by_year'].items())[:5]:
            print(f"  {year}: {count}")

        print(f"\nTop authors:")
        for author, count in list(stats['by_author'].items())[:10]:
            print(f"  {author}: {count}")

    return 0


if __name__ == '__main__':
    exit(main())
