#!/usr/bin/env python3
"""
Embed critique comments into manuscript .docx file.

Parses critique reports and embeds them as Word comments at relevant locations
in the manuscript, making it easier for authors to address feedback directly.

Usage:
    python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1
    python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1 --version 2
    python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1 --unresolved-only

Requirements:
    pip install python-docx lxml
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Pt
    import subprocess
    from lxml import etree
except ImportError:
    print("ERROR: Missing required packages. Install with:")
    print("  pip install python-docx lxml")
    sys.exit(1)


@dataclass
class CritiqueIssue:
    """Represents a single critique issue."""
    severity: str  # 'major' or 'minor'
    category: str
    description: str
    action: str
    impact: Optional[str] = None
    text_snippet: Optional[str] = None  # Text fragment to search for
    section: Optional[str] = None
    resolved: bool = False


class WordCommentHelper:
    """Helper class to add real Word comments to docx files."""

    def __init__(self, doc: Document):
        self.doc = doc
        self.comment_id = 0
        self._ensure_comments_part()

    def _ensure_comments_part(self):
        """Ensure the document has a comments part."""
        # Get the document's package
        package = self.doc.part.package

        # Check if comments part exists
        try:
            self.comments_part = package.part_related_by(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
            )
        except KeyError:
            # Comments part doesn't exist, would need to create it
            # For now, we'll use a workaround with highlighted text
            self.comments_part = None

    def add_comment(self, paragraph, text: str, comment_text: str, author: str = "RRWrite Critique") -> bool:
        """
        Add a comment to specific text in a paragraph.

        Since python-docx doesn't support comments natively, we'll use highlighting
        and add the comment text nearby.

        Args:
            paragraph: The paragraph object
            text: The text to highlight
            comment_text: The comment content
            author: Comment author

        Returns:
            True if comment was added
        """
        # For now, use a workaround: highlight text and add comment as styled paragraph
        # In the future, could implement full OOXML comment support

        para_text = paragraph.text
        if text.lower() not in para_text.lower():
            return False

        # Find the text position
        start_idx = para_text.lower().find(text.lower())
        if start_idx == -1:
            return False

        # Create a new paragraph after this one for the comment
        # (Proper implementation would insert comment range in XML)
        comment_para = paragraph.insert_paragraph_before()
        comment_para.style = 'Normal'

        # Add comment marker
        marker = comment_para.add_run(f"💬 COMMENT on \"{text[:50]}...\"")
        marker.font.bold = True
        marker.font.color.rgb = RGBColor(200, 0, 0)

        # Add comment content
        content = comment_para.add_run(f"\n{comment_text}")
        content.font.color.rgb = RGBColor(150, 0, 0)
        content.font.italic = True
        content.font.size = Pt(10)

        # Add author and date
        meta = comment_para.add_run(f"\n— {author}, {datetime.now().strftime('%Y-%m-%d')}")
        meta.font.size = Pt(9)
        meta.font.color.rgb = RGBColor(100, 100, 100)

        self.comment_id += 1
        return True


class CritiqueCommentEmbedder:
    """Embeds critique comments into Word document."""

    def __init__(self, manuscript_dir: Path, version: int = 1):
        """
        Initialize embedder.

        Args:
            manuscript_dir: Path to manuscript directory
            version: Critique version to embed (default: latest)
        """
        self.manuscript_dir = Path(manuscript_dir).resolve()
        self.version = version
        self.issues: List[CritiqueIssue] = []

    def parse_critique_reports(self) -> List[CritiqueIssue]:
        """Parse critique reports and extract issues."""
        issues = []

        # Parse content critique
        content_file = self.manuscript_dir / f"critique_content_v{self.version}.md"
        if content_file.exists():
            issues.extend(self._parse_critique_file(content_file, "content"))

        # Parse format critique
        format_file = self.manuscript_dir / f"critique_format_v{self.version}.md"
        if format_file.exists():
            issues.extend(self._parse_critique_file(format_file, "format"))

        self.issues = issues
        return issues

    def _parse_critique_file(self, file_path: Path, critique_type: str) -> List[CritiqueIssue]:
        """Parse a single critique report file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        issues = []

        # Parse major issues section
        major_section = re.search(
            r'## Major Issues.*?\n(.*?)(?=##|\Z)',
            content,
            re.DOTALL | re.IGNORECASE
        )
        if major_section:
            issues.extend(self._extract_issues(major_section.group(1), "major", critique_type))

        # Parse minor issues section
        minor_section = re.search(
            r'## Minor Issues.*?\n(.*?)(?=##|\Z)',
            content,
            re.DOTALL | re.IGNORECASE
        )
        if minor_section:
            issues.extend(self._extract_issues(minor_section.group(1), "minor", critique_type))

        # Parse warnings (treat as minor)
        warnings_section = re.search(
            r'## Warnings.*?\n(.*?)(?=##|\Z)',
            content,
            re.DOTALL | re.IGNORECASE
        )
        if warnings_section:
            issues.extend(self._extract_issues(warnings_section.group(1), "minor", critique_type))

        return issues

    def _extract_issues(self, section_text: str, severity: str, critique_type: str) -> List[CritiqueIssue]:
        """Extract individual issues from a section."""
        issues = []

        # Pattern: 1. **Category:** Description\n   - **Impact:** ...\n   - **Action:** ...
        pattern = r'\d+\.\s+\*\*([^:]+):\*\*\s+([^\n]+)(?:\n\s+- \*\*Impact:\*\*\s+([^\n]+))?(?:\n\s+- \*\*Action:\*\*\s+([^\n]+))?'

        for match in re.finditer(pattern, section_text):
            category = match.group(1).strip()
            description = match.group(2).strip()
            impact = match.group(3).strip() if match.group(3) else None
            action = match.group(4).strip() if match.group(4) else "Review and address"

            # Try to extract text snippet from description
            text_snippet = self._extract_text_snippet(description)

            # Infer section from category or description
            section = self._infer_section(category, description)

            issues.append(CritiqueIssue(
                severity=severity,
                category=f"{critique_type.title()}: {category}",
                description=description,
                action=action,
                impact=impact,
                text_snippet=text_snippet,
                section=section
            ))

        return issues

    def _extract_text_snippet(self, description: str) -> Optional[str]:
        """Extract quoted text snippet from description."""
        # Look for quoted text
        quote_match = re.search(r'"([^"]{10,100})"', description)
        if quote_match:
            return quote_match.group(1)

        # Look for text after ":" or "claim:"
        text_match = re.search(r'(?:claim|text|sentence):\s*"?([^"]{10,100})"?', description, re.IGNORECASE)
        if text_match:
            return text_match.group(1)

        return None

    def _infer_section(self, category: str, description: str) -> Optional[str]:
        """Infer manuscript section from category or description."""
        text = f"{category} {description}".lower()

        section_keywords = {
            'abstract': ['abstract'],
            'introduction': ['introduction', 'research question'],
            'methods': ['methods', 'reproducibility', 'parameter'],
            'results': ['results', 'interpretation'],
            'discussion': ['discussion'],
            'references': ['citation', 'reference', 'bibliography']
        }

        for section, keywords in section_keywords.items():
            if any(kw in text for kw in keywords):
                return section

        return None

    def generate_docx_with_comments(
        self,
        output_file: Optional[Path] = None,
        unresolved_only: bool = False
    ) -> Path:
        """
        Generate .docx file with embedded critique comments.

        Args:
            output_file: Output path (default: manuscript_with_comments_v{version}.docx)
            unresolved_only: Only include unresolved issues

        Returns:
            Path to generated .docx file
        """
        if output_file is None:
            output_file = self.manuscript_dir / f"manuscript_with_comments_v{self.version}.docx"

        # First convert markdown to docx using pandoc
        markdown_file = self.manuscript_dir / "manuscript_full.md"
        if not markdown_file.exists():
            raise FileNotFoundError(f"Manuscript not found: {markdown_file}")

        print(f"Converting markdown to .docx...")
        temp_docx = self.manuscript_dir / "manuscript_temp.docx"
        self._convert_markdown_to_docx(markdown_file, temp_docx)

        # Load the docx
        print(f"Loading document...")
        doc = Document(str(temp_docx))

        # Filter issues
        issues_to_embed = self.issues
        if unresolved_only:
            issues_to_embed = [i for i in self.issues if not i.resolved]

        # Add comments
        print(f"Embedding {len(issues_to_embed)} critique comments...")
        comments_added = 0

        for issue in issues_to_embed:
            # Try to find location in document
            if issue.text_snippet:
                # Search for specific text
                if self._add_comment_to_text(doc, issue):
                    comments_added += 1
                else:
                    # Fallback: add to section header or end
                    if self._add_comment_to_section(doc, issue):
                        comments_added += 1
            else:
                # No specific text, add to section
                if self._add_comment_to_section(doc, issue):
                    comments_added += 1

        # If no comments were added to specific locations, add summary at end
        if comments_added == 0 and issues_to_embed:
            self._add_summary_section(doc, issues_to_embed)

        # Save
        doc.save(str(output_file))

        # Clean up temp file
        if temp_docx.exists():
            temp_docx.unlink()

        print(f"\n✓ Generated: {output_file}")
        print(f"  Comments added: {comments_added}")
        print(f"  Total issues: {len(issues_to_embed)}")

        return output_file

    def _convert_markdown_to_docx(self, markdown_file: Path, output_file: Path):
        """Convert markdown to docx using pandoc."""
        cmd = [
            'pandoc',
            str(markdown_file),
            '-o', str(output_file),
            '-f', 'markdown',
            '-t', 'docx',
            '--standalone'
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"Pandoc conversion failed: {result.stderr}")

    def _add_comment_to_text(self, doc: Document, issue: CritiqueIssue) -> bool:
        """
        Add comment to specific text in document.

        Uses WordCommentHelper to add properly formatted comments.
        """
        comment_helper = WordCommentHelper(doc)
        comment_text = self._format_comment(issue)

        # Search for text in paragraphs
        for para in doc.paragraphs:
            if issue.text_snippet and issue.text_snippet.lower() in para.text.lower():
                # Found the text, add comment
                if comment_helper.add_comment(para, issue.text_snippet, comment_text):
                    return True

        return False

    def _add_comment_to_section(self, doc: Document, issue: CritiqueIssue) -> bool:
        """Add comment to section header."""
        if not issue.section:
            return False

        comment_helper = WordCommentHelper(doc)
        comment_text = self._format_comment(issue)

        # Find section header
        section_pattern = re.compile(rf'^#*\s*{re.escape(issue.section)}', re.IGNORECASE)

        for para in doc.paragraphs:
            if section_pattern.match(para.text):
                # Add comment to section header
                if comment_helper.add_comment(para, para.text[:20], comment_text):
                    return True

        return False

    def _add_summary_section(self, doc: Document, issues: List[CritiqueIssue]):
        """Add summary section with all comments at the end."""
        # Add page break
        doc.add_page_break()

        # Add header
        header = doc.add_heading('Critique Comments', 1)

        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Critique version: {self.version}")
        doc.add_paragraph(f"Total issues: {len(issues)}")

        # Group by severity
        major_issues = [i for i in issues if i.severity == "major"]
        minor_issues = [i for i in issues if i.severity == "minor"]

        if major_issues:
            doc.add_heading('Major Issues', 2)
            for i, issue in enumerate(major_issues, 1):
                para = doc.add_paragraph(f"{i}. ")
                para.add_run(f"[{issue.category}] ").bold = True
                para.add_run(issue.description)

                if issue.impact:
                    detail = para.add_run(f"\nImpact: {issue.impact}")
                    detail.italic = True

                action = para.add_run(f"\nAction: {issue.action}")
                action.font.color.rgb = RGBColor(0, 0, 255)  # Blue

        if minor_issues:
            doc.add_heading('Minor Issues & Warnings', 2)
            for i, issue in enumerate(minor_issues, 1):
                para = doc.add_paragraph(f"{i}. ")
                para.add_run(f"[{issue.category}] ").bold = True
                para.add_run(issue.description)

                action = para.add_run(f"\nAction: {issue.action}")
                action.font.color.rgb = RGBColor(0, 0, 255)  # Blue

    def _format_comment(self, issue: CritiqueIssue) -> str:
        """Format comment text."""
        parts = [
            f"[{issue.severity.upper()}]",
            f"{issue.category}:",
            issue.description
        ]

        if issue.impact:
            parts.append(f"Impact: {issue.impact}")

        parts.append(f"Action: {issue.action}")

        return " | ".join(parts)


def main():
    parser = argparse.ArgumentParser(
        description="Embed critique comments into manuscript .docx file",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Embed latest critique comments
  python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1

  # Embed specific version
  python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1 --version 2

  # Only unresolved issues
  python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1 --unresolved-only

  # Custom output file
  python scripts/rrwrite-embed-critique-comments.py --manuscript-dir manuscript/repo_v1 --output my_manuscript.docx
        """
    )

    parser.add_argument(
        '--manuscript-dir',
        required=True,
        help='Manuscript directory containing critique reports'
    )
    parser.add_argument(
        '--version',
        type=int,
        default=None,
        help='Critique version to embed (default: latest)'
    )
    parser.add_argument(
        '--output',
        help='Output .docx file path (default: manuscript_with_comments_v{version}.docx)'
    )
    parser.add_argument(
        '--unresolved-only',
        action='store_true',
        help='Only include unresolved issues'
    )

    args = parser.parse_args()

    # Validate manuscript directory
    manuscript_dir = Path(args.manuscript_dir)
    if not manuscript_dir.exists():
        print(f"✗ Manuscript directory not found: {manuscript_dir}")
        sys.exit(1)

    # Determine version
    version = args.version
    if version is None:
        # Find latest version
        critique_files = list(manuscript_dir.glob("critique_*_v*.md"))
        if not critique_files:
            print(f"✗ No critique reports found in {manuscript_dir}")
            sys.exit(1)

        versions = []
        for f in critique_files:
            match = re.search(r'_v(\d+)\.md$', f.name)
            if match:
                versions.append(int(match.group(1)))

        version = max(versions) if versions else 1

    print(f"Using critique version: {version}")

    # Check for critique files
    content_file = manuscript_dir / f"critique_content_v{version}.md"
    format_file = manuscript_dir / f"critique_format_v{version}.md"

    if not content_file.exists() and not format_file.exists():
        print(f"✗ No critique reports found for version {version}")
        sys.exit(1)

    # Create embedder
    embedder = CritiqueCommentEmbedder(manuscript_dir, version)

    # Parse critiques
    print(f"Parsing critique reports...")
    issues = embedder.parse_critique_reports()
    print(f"Found {len(issues)} issues")

    # Generate docx with comments
    output_file = Path(args.output) if args.output else None

    try:
        result_file = embedder.generate_docx_with_comments(
            output_file=output_file,
            unresolved_only=args.unresolved_only
        )

        print(f"\nNext steps:")
        print(f"1. Open in Microsoft Word: {result_file}")
        print(f"2. Review highlighted comments")
        print(f"3. Address each issue and mark as resolved")

    except Exception as e:
        print(f"\n✗ Failed to generate document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
